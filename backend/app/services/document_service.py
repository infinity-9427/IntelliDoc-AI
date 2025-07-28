import io
import logging
from pathlib import Path
from typing import Dict, Any, Optional
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

logger = logging.getLogger(__name__)

class DocumentGenerator:
    """Service for generating documents from processed results"""
    
    @staticmethod
    def create_docx_from_results(result: Dict[str, Any]) -> io.BytesIO:
        """
        Create a DOCX document from processing results
        """
        try:
            # Create new document
            doc = Document()
            
            # Add document title
            title = doc.add_heading('Document Processing Results', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add metadata section
            metadata_para = doc.add_paragraph()
            metadata_para.add_run('Processing Summary\n').bold = True
            metadata_para.add_run(f"Filename: {result.get('filename', 'Unknown')}\n")
            metadata_para.add_run(f"File Type: {result.get('file_type', 'Unknown')}\n")
            metadata_para.add_run(f"Processing Date: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            metadata_para.add_run(f"Processing Time: {result.get('processing_time', 0):.2f} seconds\n")
            metadata_para.add_run(f"Text Confidence: {(result.get('text_confidence', 0) * 100):.1f}%\n")
            
            if result.get('page_count'):
                metadata_para.add_run(f"Pages: {result.get('page_count')}\n")
            
            if result.get('text_statistics'):
                stats = result['text_statistics']
                metadata_para.add_run(f"Word Count: {stats.get('word_count', 0)}\n")
                metadata_para.add_run(f"Character Count: {stats.get('character_count', 0)}\n")
            
            # Add separator
            doc.add_paragraph('─' * 50)
            
            # Add AI Analysis if available
            if result.get('document_classification'):
                doc.add_heading('Document Classification', level=1)
                classification = result['document_classification']
                class_para = doc.add_paragraph()
                class_para.add_run('Type: ').bold = True
                class_para.add_run(f"{classification.get('type', 'Unknown').title()}\n")
                class_para.add_run('Confidence: ').bold = True
                class_para.add_run(f"{(classification.get('confidence', 0) * 100):.1f}%\n")
            
            # Add entities if available
            if result.get('entities') and len(result['entities']) > 0:
                doc.add_heading('Extracted Entities', level=1)
                entities_para = doc.add_paragraph()
                
                for entity in result['entities'][:20]:  # Limit to top 20
                    entities_para.add_run(f"• {entity.get('text', '')}").bold = True
                    entities_para.add_run(f" ({entity.get('label', 'Unknown')}) ")
                    entities_para.add_run(f"[{(entity.get('confidence', 0) * 100):.1f}%]\n")
            
            # Add sentiment analysis
            if result.get('sentiment_analysis'):
                doc.add_heading('Sentiment Analysis', level=1)
                sentiment = result['sentiment_analysis']
                sentiment_para = doc.add_paragraph()
                sentiment_para.add_run('Overall Sentiment: ').bold = True
                sentiment_para.add_run(f"{sentiment.get('overall_sentiment', 'neutral').title()}\n")
                sentiment_para.add_run('Confidence: ').bold = True
                sentiment_para.add_run(f"{(sentiment.get('confidence', 0) * 100):.1f}%\n")
                
                if sentiment.get('polarity') is not None:
                    sentiment_para.add_run('Polarity: ').bold = True
                    sentiment_para.add_run(f"{sentiment.get('polarity', 0):.2f}\n")
                
                if sentiment.get('subjectivity') is not None:
                    sentiment_para.add_run('Subjectivity: ').bold = True
                    sentiment_para.add_run(f"{sentiment.get('subjectivity', 0):.2f}\n")
            
            # Add key information
            if result.get('key_information'):
                doc.add_heading('Key Information', level=1)
                key_info = result['key_information']
                
                for info_type, values in key_info.items():
                    if values and len(values) > 0:
                        info_para = doc.add_paragraph()
                        info_para.add_run(f"{info_type.replace('_', ' ').title()}:\n").bold = True
                        for value in values[:10]:  # Limit to 10 items per type
                            info_para.add_run(f"  • {value}\n")
            
            # Add summary
            if result.get('summary'):
                doc.add_heading('Summary', level=1)
                summary_para = doc.add_paragraph(result['summary'])
                summary_para.style.font.size = Pt(11)
            
            # Add extracted text
            if result.get('extracted_text'):
                doc.add_page_break()
                doc.add_heading('Extracted Text', level=1)
                
                # Split text into paragraphs
                text_content = result['extracted_text']
                paragraphs = text_content.split('\n\n')
                
                for para_text in paragraphs:
                    if para_text.strip():
                        text_para = doc.add_paragraph(para_text.strip())
                        text_para.style.font.size = Pt(10)
            
            # Save to BytesIO
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            logger.info(f"Successfully created DOCX document")
            return doc_buffer
            
        except Exception as e:
            logger.error(f"Error creating DOCX document: {str(e)}")
            # Create error document
            error_doc = Document()
            error_doc.add_heading('Processing Error', 0)
            error_para = error_doc.add_paragraph()
            error_para.add_run('An error occurred while creating the document: ').bold = True
            error_para.add_run(str(e))
            
            error_buffer = io.BytesIO()
            error_doc.save(error_buffer)
            error_buffer.seek(0)
            return error_buffer
    
    @staticmethod
    def create_text_from_results(result: Dict[str, Any]) -> str:
        """
        Create a plain text document from processing results
        """
        try:
            lines = []
            lines.append("=== DOCUMENT PROCESSING RESULTS ===")
            lines.append("")
            
            # Add metadata
            lines.append("PROCESSING SUMMARY:")
            lines.append(f"Filename: {result.get('filename', 'Unknown')}")
            lines.append(f"File Type: {result.get('file_type', 'Unknown')}")
            lines.append(f"Processing Date: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            lines.append(f"Processing Time: {result.get('processing_time', 0):.2f} seconds")
            lines.append(f"Text Confidence: {(result.get('text_confidence', 0) * 100):.1f}%")
            
            if result.get('page_count'):
                lines.append(f"Pages: {result.get('page_count')}")
            
            if result.get('text_statistics'):
                stats = result['text_statistics']
                lines.append(f"Word Count: {stats.get('word_count', 0)}")
                lines.append(f"Character Count: {stats.get('character_count', 0)}")
            
            lines.append("")
            lines.append("=" * 50)
            lines.append("")
            
            # Add AI Analysis
            if result.get('document_classification'):
                lines.append("DOCUMENT CLASSIFICATION:")
                classification = result['document_classification']
                lines.append(f"Type: {classification.get('type', 'Unknown').title()}")
                lines.append(f"Confidence: {(classification.get('confidence', 0) * 100):.1f}%")
                lines.append("")
            
            # Add entities
            if result.get('entities') and len(result['entities']) > 0:
                lines.append("EXTRACTED ENTITIES:")
                for entity in result['entities'][:20]:
                    lines.append(f"• {entity.get('text', '')} ({entity.get('label', 'Unknown')}) [{(entity.get('confidence', 0) * 100):.1f}%]")
                lines.append("")
            
            # Add sentiment
            if result.get('sentiment_analysis'):
                lines.append("SENTIMENT ANALYSIS:")
                sentiment = result['sentiment_analysis']
                lines.append(f"Overall Sentiment: {sentiment.get('overall_sentiment', 'neutral').title()}")
                lines.append(f"Confidence: {(sentiment.get('confidence', 0) * 100):.1f}%")
                if sentiment.get('polarity') is not None:
                    lines.append(f"Polarity: {sentiment.get('polarity', 0):.2f}")
                if sentiment.get('subjectivity') is not None:
                    lines.append(f"Subjectivity: {sentiment.get('subjectivity', 0):.2f}")
                lines.append("")
            
            # Add key information
            if result.get('key_information'):
                lines.append("KEY INFORMATION:")
                key_info = result['key_information']
                for info_type, values in key_info.items():
                    if values and len(values) > 0:
                        lines.append(f"{info_type.replace('_', ' ').title()}:")
                        for value in values[:10]:
                            lines.append(f"  • {value}")
                        lines.append("")
            
            # Add summary
            if result.get('summary'):
                lines.append("SUMMARY:")
                lines.append(result['summary'])
                lines.append("")
            
            # Add extracted text
            if result.get('extracted_text'):
                lines.append("=" * 50)
                lines.append("EXTRACTED TEXT:")
                lines.append("=" * 50)
                lines.append("")
                
                # Clean and format the extracted text for better readability
                formatted_text = DocumentGenerator._format_extracted_text(result['extracted_text'])
                lines.append(formatted_text)
            
            return '\n'.join(lines)
            
        except Exception as e:
            logger.error(f"Error creating text document: {str(e)}")
            return f"Error creating text document: {str(e)}"
    
    @staticmethod
    def _format_extracted_text(extracted_text: str) -> str:
        """
        Format extracted text for better readability by cleaning up OCR artifacts
        """
        import re
        
        if not extracted_text:
            return ""
        
        # Clean up the text
        formatted_text = extracted_text
        
        # Remove excessive whitespace and fix spacing
        formatted_text = re.sub(r'\s+', ' ', formatted_text)
        
        # Fix page break markers
        formatted_text = re.sub(r'\n--- Page \d+ ---\n', '\n\n--- PAGE BREAK ---\n\n', formatted_text)
        
        # Fix OCR supplementary sections
        formatted_text = re.sub(r'\n--- OCR.*?---\n', '\n\n--- OCR SECTION ---\n', formatted_text)
        
        # Improve paragraph structure
        # Split on logical breaks and clean each section
        sections = formatted_text.split('\n\n')
        cleaned_sections = []
        
        for section in sections:
            # Clean up each section
            section = section.strip()
            if section:
                # Remove excessive spaces within lines
                section = re.sub(r' {2,}', ' ', section)
                
                # Fix common OCR spacing issues
                section = re.sub(r'\s+([,.;:!?])', r'\1', section)  # Remove spaces before punctuation
                section = re.sub(r'([.!?])([A-Z])', r'\1 \2', section)  # Add space after sentence endings
                
                # Fix phone number formatting
                section = re.sub(r'\{\s*(\d{3})-?(\d{3})-?(\d{4})\s*\}', r'(\1) \2-\3', section)
                
                # Fix email formatting
                section = re.sub(r'\{\s*([^@\s]+@[^@\s]+\.[^@\s]+)\s*\}', r'\1', section)
                
                # Fix website formatting
                section = re.sub(r'\{\s*([a-zA-Z0-9.-]+\.[a-zA-Z]{2,})\s*\}', r'\1', section)
                
                # Fix social media formatting
                section = re.sub(r'\{([^}]+)\}', r'\1', section)
                
                # Remove excessive newlines within sections
                section = re.sub(r'\n{3,}', '\n\n', section)
                
                cleaned_sections.append(section)
        
        # Rejoin sections with appropriate spacing
        formatted_text = '\n\n'.join(cleaned_sections)
        
        # Final cleanup
        formatted_text = re.sub(r'\n{3,}', '\n\n', formatted_text)
        formatted_text = formatted_text.strip()
        
        return formatted_text
